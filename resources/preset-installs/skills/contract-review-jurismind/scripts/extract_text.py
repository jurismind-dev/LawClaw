"""
提取 .docx 文件的全部文本内容（用于预分析阶段）。
用法: python extract_text.py <docx文件路径>
输出: 将全文文本打印到 stdout。
"""
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("错误: 需要安装 python-docx (pip install python-docx)", file=sys.stderr)
    sys.exit(1)


def extract(docx_path: str) -> str:
    doc = Document(docx_path)
    lines = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            lines.append(text)

    # 同时提取表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


def main():
    if len(sys.argv) < 2:
        print("用法: python extract_text.py <docx文件路径>", file=sys.stderr)
        sys.exit(1)

    docx_path = Path(sys.argv[1]).resolve()
    if not docx_path.exists():
        print(f"错误: 文件不存在 {docx_path}", file=sys.stderr)
        sys.exit(1)

    text = extract(str(docx_path))
    print(text)


if __name__ == "__main__":
    main()
